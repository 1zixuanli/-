# -*- coding: utf-8 -*-
"""根据中期检查内容生成 Word 报告"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_run_font(run, name="宋体", size=12, bold=False):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), name)


def add_heading_center(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, "黑体", size, bold=True)


def add_body(doc, text, first_indent=True, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, "宋体", size)
    if first_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5


def add_section_title(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, "黑体", 14, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def main():
    root = Path(__file__).resolve().parents[1]
    # 英文文件名便于在各环境下定位；内容与中文报告一致
    out = root / "Midterm-Inspection-Report.docx"

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    add_heading_center(doc, "毕业论文（设计）中期检查报告", 22)
    add_heading_center(doc, "（文本总结稿）", 12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "论文题目：基于 SSM 框架的车票售卖系统设计与实现\n"
        "（学院、姓名、学号、班级、指导教师、检查日期等请按学校表格手写或另页填写）"
    )
    set_run_font(r, "宋体", 10.5)

    doc.add_paragraph()

    add_body(
        doc,
        "本报告依据《大学本科生毕业论文（设计）中期检查表》要求，对课题进展情况作系统总结，"
        "涵盖工作态度、已完成与待完成任务、存在问题、拟采取措施及论文进度等方面，便于导师评阅与学院归档。",
        first_indent=True,
    )

    sections = [
        (
            "一、工作态度",
            "本人对待毕业设计态度端正，能够主动查阅资料、按指导教师意见修改方案；坚持每周（或每两周）与导师沟通进度，"
            "对开题所定“SSM + Vue 前后端分离”的技术路线认真落实。目前已完成前端主要功能开发与联调演示环境搭建，"
            "后续将集中精力完成服务端与论文撰写。",
        ),
        (
            "二、已完成任务（对照开题报告）",
            None,
        ),
    ]

    add_section_title(doc, sections[0][0])
    add_body(doc, sections[0][1])

    add_section_title(doc, "二、已完成任务（对照开题报告）")
    items = [
        "在指导教师指导下完成开题，明确研究目标：构建车票查询、余票校验与预订、订单管理以及管理端维护与统计等核心功能；确定采用 SSM 作为后端、Vue 作为前端的总体架构。",
        "完成与课题相关的文献与资料调研，梳理铁路/高铁客运售票业务特点及常见 B/S 架构实现方式，为需求分析与系统设计提供依据。",
        "完成系统需求分析与概要设计：划分用户模块、车票查询模块、订票与订单模块、管理端模块；拟定 REST 风格接口及统一响应格式（如 code、message、data），为前后端协作与论文“系统设计”章节提供素材。",
        "完成数据库概念结构与逻辑结构设计（用户、车次/时刻、订单等实体及关联），并形成与开题一致的表结构设想，为 MyBatis 映射与后续建表做准备。",
        "完成前端开发环境与工程化配置：建立基于 Vite + Vue 3 的工程，集成 Vue Router、Element Plus、Axios、Vuex；配置 package.json 与入口页面，可稳定执行 npm install、npm run dev 进行演示。",
        "完成用户端主要页面与流程（当前以 Mock 模拟后端）：注册与登录；首页按出发/到达站、日期、票种（含学生票、团体票）、高铁筛选查询；车次列表展示二等座、一等座、商务座、无座等票型及余票；无余票时弹窗提示不可预订；确认订单支持选择席别与购票数量、乘车人信息、模拟支付截止时间；“我的订单”支持支付、取消、退票（余票回退 Mock）、确认出行等状态流转；个人中心支持基础信息维护。",
        "完成管理端主要界面（Mock）：数据概览与按车次售票统计、用户列表、全部订单查询与筛选、车次信息新增/删除（内置车次隐藏策略）、各席别余票与放票调整；通过角色区分管理员与普通用户。",
        "已整理与后端对接的接口约定，后续 Spring MVC 按相同路径与数据格式返回即可替换前端 Mock，降低联调成本。",
    ]
    for i, t in enumerate(items, 1):
        add_body(doc, f"{i}. {t}")

    add_section_title(doc, "三、待完成任务")
    todo = [
        "在 MySQL 中完成建库、建表及必要索引与约束，编写 MyBatis Mapper 与 XML（或注解），实现用户、车次、票种/余票、订单等持久化操作。",
        "基于 SSM 搭建后端工程，实现注册登录、车票查询、余票校验、下单与订单状态变更、管理端维护与统计等接口，在 Service 层用事务保障一致性（如防超卖）。",
        "关闭前端 Mock，完成前后端联调与异常处理；视开题要求补充 Redis 等对热点查询的优化说明或实现。",
        "开展功能测试、接口测试及必要的性能或兼容性说明，整理结果写入论文。",
        "撰写并修改毕业论文各章节，统一格式、图表与参考文献。",
        "根据导师意见多轮修改，完成查重与答辩准备。",
    ]
    for i, t in enumerate(todo, 1):
        add_body(doc, f"{i}. {t}")

    add_section_title(doc, "四、现存在的问题")
    probs = [
        "后端 SSM 与数据库尚在推进，订单与余票等最终以服务端事务与锁机制为准；当前 Mock 仅用于流程演示，与生产级并发场景有差距。",
        "论文正文深度与截图、用例、数据库说明需随开发同步加强。",
        "开发环境（Node、构建配置等）曾遇依赖类问题，联调阶段仍可能暴露新环境配置问题，需预留调试时间。",
        "参考文献广度与规范性需在定稿前统一核对。",
    ]
    for i, t in enumerate(probs, 1):
        add_body(doc, f"{i}. {t}")

    add_section_title(doc, "五、拟采取的措施")
    measures = [
        "制定后端开发优先级：用户与车次查询 → 订单与支付/取消/退票 → 管理端统计与维护；分模块与前端联调。",
        "维护接口清单（路径、方法、请求/响应示例），与 Axios 调用保持一致。",
        "与导师保持固定沟通节奏，及时提交阶段性成果并根据反馈调整。",
        "论文与编码并行，每完成一模块即补充章节与截图。",
        "预留时间用于润色、格式检测、查重与答辩演练。",
    ]
    for i, t in enumerate(measures, 1):
        add_body(doc, f"{i}. {t}")

    add_section_title(doc, "六、设计（论文）进度说明")
    add_body(
        doc,
        "目前已完成：开题与文献调研、需求与概要设计、数据库设计草案、Vue 前端主要功能与 Mock 演示（含用户端、管理端、Vuex 与本地持久化）。"
        "当前阶段：SSM 后端与 MySQL 实现、前后端真实联调。"
        "计划后续：系统测试、论文初稿定稿、格式与查重、答辩准备。",
    )

    doc.add_paragraph()
    add_body(
        doc,
        "说明：本 Word 文档为中期检查内容的文字总结，与纸质《中期检查表》对应栏目一致时可直接誊写；"
        "表中“质量评价”“建议结果”等由指导教师填写。",
        first_indent=True,
        size=10.5,
    )

    doc.save(str(out))
    print("OK:", out)


if __name__ == "__main__":
    main()
